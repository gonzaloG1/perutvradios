from forms.word.form_word_design import FormWordDesign
from tkinter import filedialog, messagebox
from tkinter.messagebox import *
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx import Document
from tkinter import Tk,Frame
import tkinter as tk
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# import re

class FormWord(FormWordDesign):

    def __init__(self,master):
        super().__init__(master)
        # Atributos para las rutas de las imágenes
        # self.path_imagei = None
        # self.path_imagef = None
        # self.path_imagec = None

    def cpm():
        pass

    def abrir_word(self):
        archivos = filedialog.askopenfilenames(title="Seleccionar Archivos", filetypes=[("Archivos Word", "*.docx")])
        for archivo in archivos:
            self.listbox.insert(tk.END, archivo)

    def insertar_imagenes(self):
        if not messagebox.askokcancel("Confirmación", "¿Deseas proceder con la inserción de imágenes?"):
            return

        for i in range(self.listbox.size()):
            archivo_docx = self.listbox.get(i)

            # Extraer el nombre después de 'Datos PCM' y convertirlo a minúsculas sin espacios
            nombre_base = os.path.basename(archivo_docx).replace("Datos PCM ", "").replace(".docx", "").strip()
            nombre_base = nombre_base.replace(" ", "").lower()  # Eliminar espacios y pasar a minúsculas

            # Posibles extensiones de las imágenes
            extensiones = ['.png', '.jpg', '.jpeg']
            logo_path = None
            firma_path = None

            # Buscar el logo y la firma en las diferentes extensiones
            for ext in extensiones:
                if os.path.exists(f"D:/Misimagenes/{nombre_base}logo{ext}"):
                    logo_path = f"D:/Misimagenes/{nombre_base}logo{ext}"
                if os.path.exists(f"D:/Misimagenes/{nombre_base}firma{ext}"):
                    firma_path = f"D:/Misimagenes/{nombre_base}firma{ext}"

            # Verificar si las imágenes existen
            if not logo_path or not firma_path:
                messagebox.showerror("Error", f"Faltan imágenes para {nombre_base}.")
                return

            # Abrir el documento
            doc = Document(archivo_docx)

            # Calcular el ancho disponible del encabezado
            section = doc.sections[0]
            available_width = section.page_width - section.left_margin - section.right_margin
            max_logo_width = available_width * 0.6  # Ajustar a un 60% del ancho disponible
            max_firma_width = available_width * 0.4  # Ajustar a un 40% del ancho disponible
            max_firma_height = Inches(1)  # Ajustar altura máxima para la firma

            # Variable para verificar si el logo ya se ha insertado
            logo_insertado = False

            # Insertar el logo en el encabezado, ocupando un tamaño reducido y centrado
            for section in doc.sections:
                header = section.header
                if not logo_insertado:  # Solo insertar si no se ha insertado antes
                    paragraph = header.paragraphs[0]
                    run = paragraph.add_run()
                    
                    # Insertar el logo y ajustar el tamaño
                    run.add_picture(logo_path, width=min(Inches(max_logo_width), Inches(3)), height=Inches(1.5))  # Ajustar tamaño
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el logo
                    logo_insertado = True  # Marcar como insertado

                # Insertar la firma en el pie de página, solo una vez
                footer_paragraph = section.footer.paragraphs[0]
                if len(footer_paragraph.runs) == 0:  # Asegurarse de que no se haya insertado antes
                    run = footer_paragraph.add_run()
                    run.add_picture(firma_path, width=min(Inches(max_firma_width), Inches(2)), height=max_firma_height)  # Ajustar tamaño

                    # Alinear la firma a la derecha
                    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Guardar el documento
            doc.save(archivo_docx)

        # Mensaje de éxito
        messagebox.showinfo("Éxito", "Las imágenes se insertaron correctamente en todos los documentos.")

